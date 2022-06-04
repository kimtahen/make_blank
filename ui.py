import time
import pandas as pd
from docx import Document;
from PyQt5.QtWidgets import *
from PyQt5.QtCore import Qt, QObject, QThread, pyqtSignal

from core import make_blank

from time import sleep

args = {}

class Worker(QObject):
    finished = pyqtSignal()

    def __init__(self, parent):
        super().__init__()
        self.parent = parent;

    def run(self):
        global args
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.style = doc.styles['Table Grid']
        header = table.rows[0].cells
        header[0].text = "문제"
        header[1].text = "답변"
        header[2].text = "정답"
        for row in args['csv_data'].itertuples():
            target_row = table.add_row().cells
            target_row[0].text = row[1]
            target_row[1].text = make_blank(row[2], not args['nj'], self.parent.percentSlider.value()*0.1)
            target_row[2].text = row[2]
        doc.save(args['output_path']); 
        self.finished.emit()

class TableWidget(QTableWidget):
    def __init__(self):
        super().__init__(1,2)
        self.setRowCount(1)
        self.setColumnCount(2)
        self.setHorizontalHeaderLabels(["문제","정답"])
        self.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.setEditTriggers(QAbstractItemView.DoubleClicked)
        self.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

    def _addRow(self):
        currentRow = self.currentRow()
        self.insertRow(currentRow+1)

    def _removeRow(self):
        if self.rowCount()>0:
            selRows = []
            for row in self.selectedIndexes():
                selRows.append(row.row())
            selRows.sort()
            selRows.reverse()
            for selRow in selRows:
                self.removeRow(selRow)

class MyApp(QWidget):
    def __init__(self):
        super().__init__()
        global args
        args['nj'] = False
        args['p'] = 0.5
        args['csv_path'] = './data.csv'
        args['output_path'] = './output.docx'

        self.initUI()
    def initUI(self):
        self.setWindowTitle('Make Blank')
        self.resize(990, 800)
        #menubar
        self.menuBar = QMenuBar(self)

        filemenu = self.menuBar.addMenu('&File')
        loadAction = QAction('열기',self)
        loadAction.setShortcut('Ctrl+L')
        loadAction.triggered.connect(self.loadEvent)
        saveAction = QAction('저장',self)
        saveAction.setShortcut('Ctrl+S')
        saveAction.triggered.connect(self.saveEvent)
        filemenu.addAction(loadAction)
        filemenu.addAction(saveAction)

        runmenu = self.menuBar.addMenu('&Run')
        exportAction = QAction('빈칸 생성 후 워드로 내보내기', self)
        exportAction.setShortcut('Ctrl+R')
        exportAction.triggered.connect(self.runEvent)
        runmenu.addAction(exportAction)
        
        self.menuBar.show()

        #tablewidget
        self.tableWidget = TableWidget()


        addRowButton = QPushButton('행추가')
        addRowButton.clicked.connect(self.tableWidget._addRow);

        removeRowButton = QPushButton('행삭제')
        removeRowButton.clicked.connect(self.tableWidget._removeRow);

        josaToggle = QCheckBox('단어 전체 빈칸 생성',self);
        josaToggle.stateChanged.connect(self.changeJosa);

        self.percentSlider = QSlider(Qt.Horizontal, self)
        self.percentSlider.setRange(0,10)
        self.percentSlider.setSingleStep(1)
        self.percentSlider.setTickPosition(QSlider.TicksBelow)
        self.percentSlider.setValue(5)

        vbox = QVBoxLayout()
        vbox.setContentsMargins(10,30,10,10)
        vbox.addWidget(self.tableWidget)
        hbox = QHBoxLayout()
        hbox.addStretch(1)
        hbox.addWidget(addRowButton)
        hbox.addWidget(removeRowButton)
        hbox.addWidget(josaToggle)
        hbox.addWidget(self.percentSlider)
        hbox.addStretch(1)
        vbox.addLayout(hbox)
        self.setLayout(vbox)
        self.show()

    
    def loadEvent(self):
        global args
        path = QFileDialog.getOpenFileName(self, 'Load file', './', 'CSV(*.csv)')
        if(path[0]==""):
            return;
        args['csv_path'] = path[0]
        args['output_path'] = path[0].replace('.csv','.docx')
        args['csv_data'] = pd.read_csv(path[0],delimiter='|', keep_default_na=False);    
        numrows = len(args['csv_data'].index)
        numcols = len(args['csv_data'].columns)
        self.tableWidget.setRowCount(numrows) 
        self.tableWidget.setColumnCount(numcols-1) 
        for row in args['csv_data'].itertuples():
            self.tableWidget.setItem(row[0],0,QTableWidgetItem(str(row[2])))
            self.tableWidget.setItem(row[0],1,QTableWidgetItem(str(row[3])))
        
    def saveEvent(self):
        global args
        numrows = self.tableWidget.rowCount(); 
        numcols = self.tableWidget.columnCount(); 
        newdf = []
        for i in range(numrows):
            tmp = []
            for j in range(numcols):
                if(self.tableWidget.item(i,j)==None):
                    tmp.append("")
                else:
                    tmp.append(self.tableWidget.item(i,j).text()) 
            newdf.append(tmp)

        args['csv_data'] = pd.DataFrame(newdf, columns=['문제', '정답']);
        args['csv_data'].to_csv(args['csv_path'], sep='|')
    
    def runEvent(self):
        QMessageBox.about(self, "실행", '파일 생성 중입니다.')
        startTime = time.time()
        self.saveEvent();
        self.thread = QThread()
        self.worker = Worker(self)
        self.worker.moveToThread(self.thread)
        self.thread.started.connect(self.worker.run)
        self.worker.finished.connect(self.thread.quit)
        self.worker.finished.connect(self.worker.deleteLater)
        self.thread.finished.connect(self.thread.deleteLater)
        self.thread.start()
        '''
        global args
        startTime = time.time()
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.style = doc.styles['Table Grid']
        header = table.rows[0].cells
        header[0].text = "문제"
        header[1].text = "답변"
        header[2].text = "정답"
        for row in args['csv_data'].itertuples():
            target_row = table.add_row().cells
            target_row[0].text = row[1]
            target_row[1].text = make_blank(row[2], not args['nj'], self.percentSlider.value()*0.1)
            target_row[2].text = row[2]

        doc.save(args['output_path']);
        '''
        def win():
            QMessageBox.about(self, "실행", 'time elapsed: '+str(time.time()-startTime)+'\n'+args['output_path'])

        self.thread.finished.connect(win)
    
        
    def changeJosa(self):
        global args
        args['nj'] = not args['nj']
